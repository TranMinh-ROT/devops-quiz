#!/usr/bin/env python3
"""
Deeper exploration:
1. File 5: investigate 101 questions - is one duplicated or off by one?
2. How are options presented? (lettered or just text?)
3. Check different answer patterns across files
4. Check for 5+ option questions more carefully
"""
import os
import re
# pyrefly: ignore [missing-import]
from docx import Document

SOURCE_DIR = "source-documents"
files = sorted([f for f in os.listdir(SOURCE_DIR) if f.endswith('.docx')])

# ---- FILE 5: List all question numbers ----
print("="*80)
print("FILE 5: Question numbers")
print("="*80)
fname = files[4]
doc = Document(os.path.join(SOURCE_DIR, fname))
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    m = re.match(r'^Question\s+(\d+)', text)
    if m:
        print(f"  [{i:4d}] Q{m.group(1)}: {text[:120]}")

# ---- FILE 1: First 3 questions in full detail ----
print("\n" + "="*80)
print("FILE 1: First 3 questions detailed")
print("="*80)
fname = files[0]
doc = Document(os.path.join(SOURCE_DIR, fname))
q_count = 0
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if re.match(r'^Question\s+\d+', text):
        q_count += 1
        if q_count > 3:
            break
    if q_count >= 1 and q_count <= 3 and text:
        print(f"  [{i:4d}] {text[:200]}")
    
# ---- Check option patterns across all files ----
print("\n" + "="*80)
print("OPTION PATTERNS (first 10 per file)")
print("="*80)
for fname in files:
    doc = Document(os.path.join(SOURCE_DIR, fname))
    print(f"\n--- {fname} ---")
    count = 0
    in_question = False
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if re.match(r'^Question\s+\d+', text):
            in_question = True
            q_text = text
        elif in_question and text and not text.startswith('→') and not text.startswith('->') and not text.startswith('Đề bài') and not text.startswith('Giải thích') and not text.startswith('CORRECT') and not text.startswith('ĐÚNG') and not text.startswith('INCORRECT') and not text.startswith('SAI') and not text.startswith('Reference') and not text.startswith('Tài liệu') and not text.startswith('via') and not text.startswith('Nguồn') and not text.startswith('Save time') and not text.startswith('Incorrect') and not text.startswith('Các lựa'):
            if count < 30:
                # This might be an option line
                if re.match(r'^[A-F][\.\)]\s', text) or (len(text) < 200 and ' - ' in text and not text.startswith('Question')):
                    pass  # counted below
        # Check for 5-option questions
        if re.match(r'^E[\.\)]', text) or (text.startswith('E.') and len(text) < 200):
            count += 1
            if count <= 5:
                print(f"  5-opt [{i}] {text[:150]}")

# ---- Show answer line patterns for each file ----
print("\n" + "="*80)
print("ANSWER LINE PATTERNS (unique, per file)")  
print("="*80)
for fname in files:
    doc = Document(os.path.join(SOURCE_DIR, fname))
    print(f"\n--- {fname} ---")
    patterns = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if '→ Đáp án' in text or '-> Đáp án' in text or '--> Đáp án' in text or text.startswith('Đáp án:'):
            patterns.append((i, text[:200]))
    # Show first 5 and unique structures
    for idx, (i, t) in enumerate(patterns[:5]):
        print(f"  [{i}] {t}")
    print(f"  Total answer lines: {len(patterns)}")

# ---- File 5: look at question structure (it's different) ----
print("\n" + "="*80)
print("FILE 5: First 2 questions in detail (different structure)")
print("="*80)
fname = files[4]
doc = Document(os.path.join(SOURCE_DIR, fname))
q_count = 0
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if re.match(r'^Question\s+\d+', text):
        q_count += 1
        if q_count > 2:
            break
    if q_count >= 1 and q_count <= 2 and text:
        print(f"  [{i:4d}] {text[:250]}")
