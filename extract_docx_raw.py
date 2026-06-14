#!/usr/bin/env python3
"""
Extract raw text from all DOCX files in source-documents/ to understand their structure.
"""
import os
from docx import Document

SOURCE_DIR = "source-documents"

files = sorted([f for f in os.listdir(SOURCE_DIR) if f.endswith('.docx')])
print(f"Found {len(files)} DOCX files:\n")

for fname in files:
    fpath = os.path.join(SOURCE_DIR, fname)
    doc = Document(fpath)
    
    print(f"{'='*80}")
    print(f"FILE: {fname}")
    print(f"Paragraphs: {len(doc.paragraphs)}")
    print(f"{'='*80}")
    
    # Print first 200 paragraphs to understand structure
    for i, para in enumerate(doc.paragraphs[:200]):
        style = para.style.name if para.style else "None"
        text = para.text.strip()
        if text:
            print(f"  [{i:4d}] ({style:20s}) | {text[:200]}")
    
    print(f"\n  ... (showing first 200 non-empty paragraphs of {len(doc.paragraphs)} total)\n")
