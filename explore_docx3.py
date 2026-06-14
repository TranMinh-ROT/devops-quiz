#!/usr/bin/env python3
"""
Check how different files structure bilingual content.
File 1-4 seem to alternate English/Vietnamese lines.
File 5 seems to combine with " / " separator.
Also check option labeling patterns.
"""
import os
import re
from docx import Document

SOURCE_DIR = "source-documents"
files = sorted([f for f in os.listdir(SOURCE_DIR) if f.endswith('.docx')])

# Check File 2 structure 
print("="*80)
print("FILE 2: First 3 questions in detail")
print("="*80)
fname = files[1]
doc = Document(os.path.join(SOURCE_DIR, fname))
q_count = 0
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if re.match(r'^Question\s+\d+', text):
        q_count += 1
        if q_count > 3:
            break
    if q_count >= 1 and q_count <= 3 and text:
        print(f"  [{i:4d}] {text[:300]}")

# Check File 3 structure
print("\n" + "="*80)
print("FILE 3: First 3 questions in detail")
print("="*80)
fname = files[2]
doc = Document(os.path.join(SOURCE_DIR, fname))
q_count = 0
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if re.match(r'^Question\s+\d+', text):
        q_count += 1
        if q_count > 3:
            break
    if q_count >= 1 and q_count <= 3 and text:
        print(f"  [{i:4d}] {text[:300]}")

# Check File 4 structure
print("\n" + "="*80)
print("FILE 4: First 3 questions in detail")
print("="*80)
fname = files[3]
doc = Document(os.path.join(SOURCE_DIR, fname))
q_count = 0
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if re.match(r'^Question\s+\d+', text):
        q_count += 1
        if q_count > 3:
            break
    if q_count >= 1 and q_count <= 3 and text:
        print(f"  [{i:4d}] {text[:300]}")

# Check File 5 question with 5 options
print("\n" + "="*80)
print("FILE 5: Check for 5-option question (near option E)")
print("="*80)
fname = files[4]
doc = Document(os.path.join(SOURCE_DIR, fname))
# Find option E
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if 'E.' in text and i > 0:
        # Print context around it
        start = max(0, i-8)
        end = min(len(doc.paragraphs), i+5)
        for j in range(start, end):
            t = doc.paragraphs[j].text.strip()
            if t:
                marker = " <<<" if j == i else ""
                print(f"  [{j:4d}] {t[:200]}{marker}")
        print("---")

# Check how the duplicate Q40 looks in file 5
print("\n" + "="*80)
print("FILE 5: Duplicate Q40 (paragraphs 606-640)")
print("="*80)
fname = files[4]
doc = Document(os.path.join(SOURCE_DIR, fname))
for i in range(606, 645):
    text = doc.paragraphs[i].text.strip()
    if text:
        print(f"  [{i:4d}] {text[:250]}")
