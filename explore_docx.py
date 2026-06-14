#!/usr/bin/env python3
"""
Deep exploration of DOCX structure - dump ALL paragraphs from all files.
Also check for multi-answer questions and count questions per file.
"""
import os
import re
from docx import Document

SOURCE_DIR = "source-documents"

files = sorted([f for f in os.listdir(SOURCE_DIR) if f.endswith('.docx')])

for fname in files:
    fpath = os.path.join(SOURCE_DIR, fname)
    doc = Document(fpath)
    
    # Count questions
    q_count = 0
    multi_answer_qs = []
    
    all_text = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            all_text.append((i, text))
            # Count question headers
            if re.match(r'^Question\s+\d+', text):
                q_count += 1
            # Find multi-answer indicators
            if re.search(r'(Select\s+(TWO|THREE|FOUR|ALL)|Choose\s+(TWO|THREE|FOUR)|Which\s+(TWO|THREE)|Chọn\s+(HAI|BA)|select\s+all\s+that\s+apply)', text, re.IGNORECASE):
                multi_answer_qs.append((i, text[:150]))
    
    print(f"\n{'='*80}")
    print(f"FILE: {fname}")
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    print(f"Non-empty paragraphs: {len(all_text)}")
    print(f"Questions found: {q_count}")
    print(f"Multi-answer questions: {len(multi_answer_qs)}")
    for idx, txt in multi_answer_qs:
        print(f"  [{idx}] {txt}")
    
    # Show how answer lines look
    print(f"\nAnswer line patterns:")
    answer_patterns = set()
    for i, text in all_text:
        if text.startswith('→ Đáp án') or text.startswith('-> Đáp án') or text.startswith('--> Đáp án') or text.startswith('Đáp án:') or text.startswith('→ Answer'):
            answer_patterns.add(text[:120])
            if len(answer_patterns) <= 10:
                print(f"  [{i}] {text[:150]}")
    
    # Check for option patterns (E, F options)
    has_e = 0
    has_f = 0
    for i, text in all_text:
        if re.match(r'^E[\.\)]\s', text) or re.match(r'^E\s*[-–]', text):
            has_e += 1
        if re.match(r'^F[\.\)]\s', text) or re.match(r'^F\s*[-–]', text):
            has_f += 1
    
    print(f"\nOption E occurrences: {has_e}")
    print(f"Option F occurrences: {has_f}")
